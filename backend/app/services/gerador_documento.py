import os
from pathlib import Path

from docxtpl import DocxTemplate

BACKEND_DIR = Path(__file__).parent.parent.parent
TEMPLATE_ETP = BACKEND_DIR / "templates" / "template_etp.docx"


def gerar_etp(conteudo: dict, destino: str) -> str:
    doc = DocxTemplate(str(TEMPLATE_ETP))
    doc.render(conteudo)
    doc.save(destino)
    return destino


def gerar_tr_provisorio(conteudo: dict, destino: str) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    aviso = doc.add_paragraph()
    aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = aviso.add_run("MINUTA PROVISÓRIA — SEM FORMATAÇÃO OFICIAL DO ES")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    titulo = doc.add_heading("TERMO DE REFERÊNCIA", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_identificacao_tr(doc, conteudo)

    SECOES_TR = [
        ("objeto",                "1. OBJETO"),
        ("fundamentacao_legal",   "2. FUNDAMENTAÇÃO LEGAL"),
        ("descricao_tecnica",     "3. DESCRIÇÃO TÉCNICA"),
        ("quantidade_unidade",    "4. QUANTIDADE E UNIDADE"),
        ("criterio_julgamento",   "5. CRITÉRIO DE JULGAMENTO"),
        ("habilitacao",           "6. HABILITAÇÃO"),
        ("prazo_execucao",        "7. PRAZO DE EXECUÇÃO"),
        ("local_entrega",         "8. LOCAL DE ENTREGA"),
        ("obrigacoes_contratada", "9. OBRIGAÇÕES DA CONTRATADA"),
        ("obrigacoes_contratante","10. OBRIGAÇÕES DA CONTRATANTE"),
        ("criterio_medicao",      "11. CRITÉRIO DE MEDIÇÃO E PAGAMENTO"),
        ("valor_estimado",        "12. VALOR ESTIMADO"),
        ("dotacao_orcamentaria",  "13. DOTAÇÃO ORÇAMENTÁRIA"),
        ("penalidades",           "14. PENALIDADES"),
        ("garantia",              "15. GARANTIA"),
        ("fiscalizacao",          "16. FISCALIZAÇÃO"),
        ("disposicoes_gerais",    "17. DISPOSIÇÕES GERAIS"),
    ]

    for campo, titulo_secao in SECOES_TR:
        doc.add_heading(titulo_secao, level=1)
        texto = conteudo.get(campo, "[A PREENCHER PELO AGENTE]")
        doc.add_paragraph(str(texto))

    pendencias = conteudo.get("pendencias", [])
    if pendencias:
        doc.add_heading("PENDÊNCIAS PARA REVISÃO", level=1)
        for p in pendencias:
            doc.add_paragraph(f"• {p}", style="List Bullet")

    doc.save(destino)
    return destino


def _add_identificacao_tr(doc, conteudo: dict):
    from docx.shared import Pt

    doc.add_heading("IDENTIFICAÇÃO", level=1)
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"

    dados = [
        ("Unidade Gestora", conteudo.get("un_gestora", "")),
        ("Responsáveis", conteudo.get("responsaveis", "")),
        ("Data de Elaboração", conteudo.get("data_elab", "")),
        ("Número do Processo", conteudo.get("numero_processo", "")),
        ("Versão", "1"),
    ]
    for i, (label, valor) in enumerate(dados):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = str(valor)

    doc.add_paragraph()
