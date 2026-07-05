"""
Transforma o modelo oficial do ES em template_etp.docx com tags Jinja2 para docxtpl.
Rodar uma vez após inspecionar o modelo.

Estrutura real do modelo:
- Parágrafos 'Nivel 01' = títulos de seção
- Parágrafos 'Nivel 2' com '...' = placeholders de conteúdo (substituir por {{ campo }})
- Tabelas 1×1 com "O que deve ser informado" = notas explicativas (remover)
- Tabela 0 (5×4) = identificação (injetar tags nas células de valor)
- Tabelas 16-17 = Anexo I riscos (injetar tags por posição)
- Tabelas 20-21 = Anexo II matriz (idem)
"""
import shutil
import unicodedata
import re
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent.parent.parent
BACKEND = Path(__file__).parent.parent
MODELO = ROOT / "Modelo de ETP - Estudo Técnico Preliminar - Governo do ES - Lei 14133 - v.1 (1).docx"
DEST = BACKEND / "templates" / "template_etp.docx"
DEST.parent.mkdir(exist_ok=True)
shutil.copy2(MODELO, DEST)


def _norm(t: str) -> str:
    nfkd = unicodedata.normalize("NFKD", t or "")
    s = "".join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", s).lower().strip()


def _set_para(p, novo: str):
    """Substitui texto do parágrafo preservando formatação do primeiro run."""
    for r in p.runs[1:]:
        r.text = ""
    if p.runs:
        p.runs[0].text = novo
    else:
        p.add_run(novo)


def _set_celula(cell, novo: str):
    """Substitui texto de uma célula mantendo o parágrafo existente."""
    for pi, para in enumerate(cell.paragraphs):
        for r in para.runs:
            r.text = ""
        if pi == 0:
            if para.runs:
                para.runs[0].text = novo
            else:
                para.add_run(novo)


def _remover_tabela(tbl):
    tbl._element.getparent().remove(tbl._element)


# ---------------------------------------------------------------------------
# Mapeamento seções → variáveis Jinja (em ordem de aparição no doc)
# ---------------------------------------------------------------------------
SECOES = [
    ("descricao da necessidade da contratacao",         "descricao_necessidade"),
    ("demonstracao da previsao da contratacao",         "previsao_pca"),
    ("requisitos da contratacao",                       "requisitos_contratacao"),
    ("estimativas das quantidades",                     "estimativa_quantidade"),
    ("levantamento de mercado",                         "levantamento_mercado"),
    ("estimativa do valor",                             "estimativa_valor"),
    ("descricao da solucao",                            "descricao_solucao"),
    ("justificativas para o parcelamento",              "justificativa_parcelamento"),
    ("demonstrativo dos resultados pretendidos",        "resultados_pretendidos"),
    ("providencias a serem adotadas",                   "providencias_previas"),
    ("contratacoes correlatas",                         "contratacoes_correlatas"),
    ("possiveis impactos ambientais",                   "impactos_ambientais"),
    ("posicionamento conclusivo",                       "posicionamento_conclusivo"),
]

# ---------------------------------------------------------------------------
# Identificação: texto da label → variável
# ---------------------------------------------------------------------------
IDENT = {
    "un. gestora":              "un_gestora",
    "un. adm. envolvidas":      "un_adm_envolvidas",
    "responsaveis":             "responsaveis",
    "data de elab":             "data_elab",
    "versao":                   "versao",
    "objeto":                   "objeto_resumido",
    "numero do processo":       "numero_processo",
    "programa de trabalho":     "programa_trabalho",
}


def _processar_identificacao(tbl):
    substituidos = 0
    for row in tbl.rows:
        cells = row.cells
        for ci, cell in enumerate(cells):
            label = _norm(cell.text)
            for chave, var in IDENT.items():
                if chave in label and ci + 1 < len(cells):
                    valor_cell = cells[ci + 1]
                    # Não sobrescreve se já tem tag
                    if not valor_cell.text.strip().startswith("{{"):
                        _set_celula(valor_cell, f"{{{{ {var} }}}}")
                        substituidos += 1
                    break
    return substituidos


# ---------------------------------------------------------------------------
# Riscos — injetar tags por posição nas tabelas existentes
# ---------------------------------------------------------------------------
# Cada bloco de risco (Tabela 16 e 17 → Anexo I; Tabela 20 e 21 → Anexo II)
# tem estrutura fixa de linhas. Mapeamos índice de linha → campo.

RISCO_MAP_AI = {
    # linha: (col_label, col_valor, campo)
    1: (None, 0, "{descricao}"),        # linha 1: célula 0 = "Descrição: ..."
    2: (0, 2, "{probabilidade}"),       # linha 2: col 2 = opções probabilidade
    3: (0, 2, "{impacto}"),             # linha 3: col 2 = opções impacto
    4: (0, 2, "{fase_impactada}"),      # linha 4: col 2 = fase impactada
    # linhas adicionais variam — pega as que tiver
}

def _injetar_risco_ai(tbl, idx: int):
    """Injeta tags em bloco de risco do Anexo I (linhas variáveis)."""
    prefixo = f"riscos[{idx}]"
    linhas = tbl.rows
    n = len(linhas)

    # Linha 0: header "RISCO N"
    if n > 0:
        _set_celula(linhas[0].cells[0], f"Risco {{{{ {prefixo}.id }}}}")

    # Linha 1: Descrição (célula mesclada — ocupa 4 colunas, pegar cells[0])
    if n > 1:
        c = linhas[1].cells[0]
        _set_celula(c, f"Descrição: {{{{ {prefixo}.descricao }}}}")

    # Linha 2: Probabilidade
    if n > 2:
        _set_celula(linhas[2].cells[2], f"{{{{ {prefixo}.probabilidade }}}}")

    # Linha 3: Impacto
    if n > 3:
        _set_celula(linhas[3].cells[2], f"{{{{ {prefixo}.impacto }}}}")

    # Linha 4: Fase Impactada
    if n > 4:
        _set_celula(linhas[4].cells[2], f"{{{{ {prefixo}.fase_impactada }}}}")

    # Linha 5: Dano Potencial
    if n > 5:
        _set_celula(linhas[5].cells[2], f"{{{{ {prefixo}.dano_potencial }}}}")

    # Linha 6: Ação Preventiva
    if n > 6:
        _set_celula(linhas[6].cells[2], f"{{{{ {prefixo}.acao_preventiva }}}}")

    # Linha 7: Contingência
    if n > 7:
        _set_celula(linhas[7].cells[2], f"{{{{ {prefixo}.contingencia }}}}")


def _injetar_risco_aii(tbl, idx: int):
    """Injeta tags em bloco de risco do Anexo II."""
    prefixo = f"matriz_riscos[{idx}]"
    linhas = tbl.rows
    n = len(linhas)

    if n > 0:
        _set_celula(linhas[0].cells[0], f"Risco {{{{ {prefixo}.id }}}}")
    if n > 1:
        _set_celula(linhas[1].cells[0], f"Descrição: {{{{ {prefixo}.descricao }}}}")
    if n > 2:
        _set_celula(linhas[2].cells[2], f"{{{{ {prefixo}.probabilidade }}}}")
    if n > 3:
        _set_celula(linhas[3].cells[2], f"{{{{ {prefixo}.impacto }}}}")
    if n > 4:
        # Linha 4 tem sub-colunas Id/Materialização
        _set_celula(linhas[4].cells[0], "Id")
        _set_celula(linhas[4].cells[1], f"{{{{ {prefixo}.materializacao }}}}")
    if n > 5:
        _set_celula(linhas[5].cells[0], f"{{{{ {prefixo}.acao_preventiva }}}}")
    if n > 6:
        _set_celula(linhas[6].cells[0], f"{{{{ {prefixo}.resp_privado }}}}")
    if n > 7:
        _set_celula(linhas[7].cells[0], f"{{{{ {prefixo}.resp_publico }}}}")


def main():
    doc = Document(str(DEST))

    # ------------------------------------------------------------------
    # 1. Processar parágrafos: Nivel 01 detecta seção; Nivel 2 recebe tag
    # ------------------------------------------------------------------
    secao_var = None
    subs_paras = 0

    for p in doc.paragraphs:
        estilo = p.style.name if p.style else ""
        texto = "".join(r.text for r in p.runs).strip()
        tnorm = _norm(texto)

        if re.match(r"nivel\s*0?1\b", estilo.lower()):
            secao_var = None
            for chave, var in SECOES:
                if chave in tnorm:
                    secao_var = var
                    break

        elif re.match(r"nivel\s*2\b", estilo.lower()):
            if texto == "..." and secao_var:
                _set_para(p, f"{{{{ {secao_var} }}}}")
                subs_paras += 1
                secao_var = None

    # ------------------------------------------------------------------
    # 2. Processar tabelas
    # ------------------------------------------------------------------
    tabelas = list(doc.tables)
    removidas = 0
    subs_ident = 0
    risco_ai_idx = 0
    risco_aii_idx = 0

    for tbl in tabelas:
        n_rows = len(tbl.rows)
        n_cols = len(tbl.columns)
        if n_rows == 0:
            continue

        first_cell = tbl.cell(0, 0).text.strip()
        fnorm = _norm(first_cell)

        # Tabela de identificação
        if "identificacao" in fnorm:
            subs_ident += _processar_identificacao(tbl)
            continue

        # Tabelas de notas (1×1 com "O que deve ser informado" ou "NOTAS EXPLICATIVAS")
        if n_rows == 1 and n_cols == 1:
            if any(k in fnorm for k in [
                "o que deve ser informado",
                "notas explicativas",
                "observacao",
                "nota:",
            ]):
                _remover_tabela(tbl)
                removidas += 1
                continue

        # Tabelas de risco Anexo I (11 linhas, 4 cols, header "RISCO N")
        if n_rows >= 8 and n_cols == 4 and "risco" in fnorm and risco_ai_idx < 5:
            # Detectar se é Anexo I (tem "Fase Impactada") ou Anexo II (tem "Id")
            linhas_texto = " ".join(c.text for c in tbl.rows[4].cells) if n_rows > 4 else ""
            if "fase" in _norm(linhas_texto):
                _injetar_risco_ai(tbl, risco_ai_idx)
                risco_ai_idx += 1
                continue
            elif "id" in _norm(linhas_texto) or "materializacao" in _norm(linhas_texto):
                _injetar_risco_aii(tbl, risco_aii_idx)
                risco_aii_idx += 1
                continue

    doc.save(str(DEST))
    print(f"[OK] {DEST.name}")
    print(f"     Parágrafos de conteúdo substituídos : {subs_paras}")
    print(f"     Células de identificação tagueadas  : {subs_ident}")
    print(f"     Tabelas de nota removidas           : {removidas}")
    print(f"     Blocos Anexo I injetados            : {risco_ai_idx}")
    print(f"     Blocos Anexo II injetados           : {risco_aii_idx}")
    print("Template pronto.")


if __name__ == "__main__":
    main()
