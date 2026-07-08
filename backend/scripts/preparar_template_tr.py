"""
Transforma a minuta PGE (20.02 TR E HABILITAÇÃO de COMPRAS.docx) em
template_tr.docx com tags Jinja2 para docxtpl.

Política por seção:
  FILL    — substituir parágrafo principal por {{ campo }}
  HYBRID  — substituir inline o placeholder ____________ pela tag
  KEEP    — não tocar (boilerplate jurídico)

Rodar uma vez; se o modelo mudar, rodar novamente.
"""
import re
import shutil
import unicodedata
from pathlib import Path

from docx import Document

ROOT = Path(__file__).parent.parent.parent
BACKEND = Path(__file__).parent.parent
MODELO = ROOT / "20.02 TR E HABILITAÇÃO de COMPRAS.docx"
DEST = BACKEND / "templates" / "template_tr.docx"
DEST.parent.mkdir(exist_ok=True)
shutil.copy2(MODELO, DEST)

# Padrão de placeholder nos parágrafos PGE
_PLACEHOLDER = re.compile(r'_{4,}|\.{6,}|\(preencher\)|\(reencher\)|\(verificar a pertinencia\)', re.IGNORECASE)
_PARENS_PLACEHOLDER = re.compile(r'\(\.{2,}.*?\)', re.DOTALL)


def _norm(t: str) -> str:
    nfkd = unicodedata.normalize("NFKD", t or "")
    s = "".join(c for c in nfkd if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", s).lower().strip()


def _set_para(p, novo: str):
    """Substitui todo o texto do parágrafo pela tag, preservando formatação do 1º run."""
    for r in p.runs[1:]:
        r.text = ""
    if p.runs:
        p.runs[0].text = novo
    else:
        p.add_run(novo)


def _remover_paragrafo(p):
    p._element.getparent().remove(p._element)


def _sub_inline(p, tag: str) -> bool:
    """Substitui o primeiro placeholder encontrado nos runs pela tag Jinja."""
    for r in p.runs:
        if _PLACEHOLDER.search(r.text) or _PARENS_PLACEHOLDER.search(r.text):
            texto_novo = _PLACEHOLDER.sub(tag, r.text)
            texto_novo = _PARENS_PLACEHOLDER.sub(tag, texto_novo)
            r.text = texto_novo
            return True
    return False


def _processar_tabela_itens(tbl):
    """Converte Tabela 0 (itens) em loop docxtpl usando 3 linhas separadas.

    docxtpl substitui a linha que contém {%tr for %} pelo bloco Jinja {% for %}
    e a linha que contém {%tr endfor %} pelo {% endfor %}.  Os dois tags precisam
    estar em linhas DISTINTAS; a linha do meio é o template repetível.

    Estrutura após o processamento:
      header
      linha A  → {%tr for it in itens %}     (vira {% for it in itens %})
      linha B  → {{ it.numero }} … etc.       (linha repetida para cada item)
      linha C  → {%tr endfor %}               (vira {% endfor %})
    """
    rows = tbl.rows
    if len(rows) < 4:
        return  # precisa de header + pelo menos 3 linhas de dados

    campos = [
        "{{ it.numero }}",
        "{{ it.descricao }}",
        "{{ it.unidade }}",
        "{{ it.quantidade }}",
        "{{ it.valor_unitario_max }}",
        "{{ it.valor_total }}",
    ]

    # Linha 1 → início do loop (só cell[0] com tag; demais vazias)
    for ci, cell in enumerate(rows[1].cells):
        _set_celula(cell, "{%tr for it in itens %}" if ci == 0 else "")

    # Linha 2 → conteúdo repetível
    for ci, cell in enumerate(rows[2].cells):
        _set_celula(cell, campos[ci] if ci < len(campos) else "")

    # Linha 3 → fim do loop
    for ci, cell in enumerate(rows[3].cells):
        _set_celula(cell, "{%tr endfor %}" if ci == 0 else "")

    # Deletar linhas extras (índice 4+)
    for row in list(tbl.rows)[4:]:
        row._element.getparent().remove(row._element)


def _set_celula(cell, novo: str):
    for pi, para in enumerate(cell.paragraphs):
        for r in para.runs:
            r.text = ""
        if pi == 0:
            if para.runs:
                para.runs[0].text = novo
            else:
                para.add_run(novo)


def main():
    doc = Document(str(DEST))
    paras = list(doc.paragraphs)

    to_remove: set[int] = set()  # id(p) para remoção posterior
    secao = None
    fills_done: set[str] = set()  # seções onde já fizemos o fill principal

    fill_count = 0
    inline_count = 0
    notas_count = 0
    alt_count = 0

    for p in paras:
        style = p.style.name if p.style else ""
        texto = "".join(r.text for r in p.runs).strip()
        tnorm = _norm(texto)

        # ── Remover todas as notas explicativas ──────────────────────────────
        if style == "PGE-NotaExplicativa":
            to_remove.add(id(p))
            notas_count += 1
            continue

        # ── Atualizar seção ao encontrar Heading 1 ───────────────────────────
        if style == "Heading 1":
            secao = tnorm
            fills_done.discard(secao)
            continue

        if secao is None:
            continue

        # ================================================================
        # FUNDAMENTAÇÃO E DESCRIÇÃO DA NECESSIDADE  — FILL
        # ================================================================
        if "fundamentacao" in secao and "necessidade" in secao:
            if style == "N 1.1" and "fundamentacao da contratacao" in tnorm:
                if secao not in fills_done:
                    _set_para(p, "{{ fundamentacao_necessidade }}")
                    fills_done.add(secao)
                    fill_count += 1
                else:
                    to_remove.add(id(p))
                    alt_count += 1
            elif tnorm == "ou" and secao in fills_done:
                to_remove.add(id(p))
                alt_count += 1

        # ================================================================
        # DESCRIÇÃO DA SOLUÇÃO — FILL
        # ================================================================
        elif "descricao da solucao" in secao:
            if style == "N 1.1" and "descricao da solucao" in tnorm:
                if secao not in fills_done:
                    _set_para(p, "{{ descricao_solucao }}")
                    fills_done.add(secao)
                    fill_count += 1
                else:
                    to_remove.add(id(p))
                    alt_count += 1
            elif tnorm == "ou" and secao in fills_done:
                to_remove.add(id(p))
                alt_count += 1

        # ================================================================
        # REQUISITOS DA CONTRATAÇÃO — FILL parcial (slots de sustentabilidade)
        # ================================================================
        elif "requisitos da contratacao" in secao:
            if style == "N 1.1.1" and _PLACEHOLDER.search(texto):
                if secao not in fills_done:
                    _set_para(p, "{{ requisitos_sustentabilidade }}")
                    fills_done.add(secao)
                    fill_count += 1
                else:
                    # Segunda ocorrência → remover (duplicata de slot)
                    to_remove.add(id(p))
                    alt_count += 1
            # Todo o resto da seção (amostras, subcontratação, garantia) → KEEP

        # ================================================================
        # MODELO DE EXECUÇÃO DO OBJETO — HYBRID
        # ================================================================
        elif "modelo de execucao" in secao:
            if not _PLACEHOLDER.search(texto):
                continue
            if "prazo de entrega" in tnorm or ("prazo" in tnorm and "dias" in tnorm and "entrega" in tnorm):
                _sub_inline(p, "{{ prazo_entrega }}")
                inline_count += 1
            elif "endereco" in tnorm or ("bens" in tnorm and "entregues" in tnorm and "seguinte" in tnorm):
                _sub_inline(p, "{{ local_entrega }}")
                inline_count += 1
            elif "validade" in tnorm and ("inferior" in tnorm or "entrega" in tnorm):
                _sub_inline(p, "{{ validade_minima }}")
                inline_count += 1
            elif "garantia contratual" in tnorm and "minimo" in tnorm:
                _sub_inline(p, "{{ garantia_meses }}")
                inline_count += 1
            elif "reparacao" in tnorm or ("prazo" in tnorm and "reparos" in tnorm):
                _sub_inline(p, "{{ prazo_reparacao }}")
                inline_count += 1

        # ================================================================
        # MODELO DE GESTÃO DO CONTRATO — KEEP + rotinas
        # ================================================================
        elif "gestao do contrato" in secao or "modelo de gestao" in secao:
            if style == "N 1.1.1" and _PLACEHOLDER.search(texto):
                _sub_inline(p, "{{ rotinas_gestao }}")
                inline_count += 1
            # Resto → KEEP

        # ================================================================
        # CRITÉRIOS DE MEDIÇÃO E DE PAGAMENTO — KEEP inteiro
        # ================================================================
        elif "criterios de medicao" in secao or ("medicao" in secao and "pagamento" in secao):
            pass

        # ================================================================
        # FORMA E CRITÉRIOS DE SELEÇÃO — HYBRID
        # ================================================================
        elif "forma e criterios" in secao or "selecao do fornecedor" in secao:
            if not _PLACEHOLDER.search(texto) and not _PARENS_PLACEHOLDER.search(texto):
                # Verificar se é placeholder de pontos (.........)
                if not re.search(r'\.{6,}', texto):
                    continue
            if "criterio" in tnorm and ("adocao do criterio" in tnorm or "julgamento" in tnorm):
                _sub_inline(p, "{{ criterio_julgamento }}")
                inline_count += 1
            elif "justificativa para adocao do referido criterio" in tnorm:
                _set_para(p, "{{ justificativa_criterio_julgamento }}")
                fill_count += 1
            elif "fornecimento do objeto sera" in tnorm:
                _sub_inline(p, "{{ forma_fornecimento }}")
                inline_count += 1
            elif "justificativa para adocao da referida forma" in tnorm:
                _set_para(p, "{{ justificativa_forma_fornecimento }}")
                fill_count += 1
            elif re.search(r'\.{6,}', texto) and style in ("N 1.1.1", "N 1.1.1.1"):
                # Slots "........." da sub-seção Das Exigências de Habilitação
                if secao not in fills_done:
                    _set_para(p, "{{ justificativa_requisitos_habilitacao }}")
                    fills_done.add(secao)
                    fill_count += 1
                else:
                    to_remove.add(id(p))
                    alt_count += 1

        # ================================================================
        # ESTIMATIVAS DO VALOR DA CONTRATAÇÃO — FILL + remover alternativas
        # ================================================================
        elif "estimativas do valor" in secao:
            if style == "N 1.1" and "custo estimado total" in tnorm:
                if secao not in fills_done:
                    _set_para(p, "{{ estimativa_valor }}")
                    fills_done.add(secao)
                    fill_count += 1
                else:
                    to_remove.add(id(p))
                    alt_count += 1
            elif tnorm == "ou" or tnorm == "ou:":
                to_remove.add(id(p))
                alt_count += 1
            elif style == "N 1.1" and secao in fills_done:
                # Alternativas (desconto, sigiloso) → remover
                to_remove.add(id(p))
                alt_count += 1

        # ================================================================
        # ADEQUAÇÃO ORÇAMENTÁRIA — KEEP + remover "Ou:" alternativa
        # ================================================================
        elif "adequacao orcamentaria" in secao:
            if tnorm in ("ou:", "ou"):
                to_remove.add(id(p))
                alt_count += 1
            # Bullets com ___ ficam para preenchimento manual

        # ================================================================
        # SANÇÕES — KEEP inteiro
        # ================================================================
        elif "sancoes administrativas" in secao:
            pass

        # ================================================================
        # HABILITAÇÃO JURÍDICA — KEEP
        # ================================================================
        elif "habilitacao juridica" in secao:
            pass

        # ================================================================
        # HABILITAÇÃO FISCAL, SOCIAL E TRABALHISTA — KEEP
        # ================================================================
        elif "fiscal" in secao and "trabalhista" in secao:
            pass

        # ================================================================
        # HABILITAÇÃO TÉCNICA — FILL pontual
        # ================================================================
        elif "habilitacao tecnica" in secao:
            if "registro ou inscricao" in tnorm and "conselho regional" in tnorm:
                _sub_inline(p, "{{ habilitacao_conselho }}")
                inline_count += 1
            elif style == "N abc" and _PARENS_PLACEHOLDER.search(texto):
                # (...descrição do objeto...) e (...indicação do quantitativo...)
                if secao not in fills_done:
                    _set_para(p, "{{ habilitacao_atestados }}")
                    fills_done.add(secao)
                    fill_count += 1
                else:
                    to_remove.add(id(p))
                    alt_count += 1
            elif "prova de atendimento aos requisitos" in tnorm:
                _set_para(p, "{{ habilitacao_lei_especifica }}")
                fill_count += 1

        # ================================================================
        # HABILITAÇÃO ECONÔMICO-FINANCEIRA — KEEP
        # ================================================================
        elif "economico" in secao and "financeira" in secao:
            pass

    # ── Executar remoções ────────────────────────────────────────────────
    for p in paras:
        if id(p) in to_remove:
            try:
                _remover_paragrafo(p)
            except Exception:
                pass

    # ── Processar Tabela 0 (loop de itens) ──────────────────────────────
    _processar_tabela_itens(doc.tables[0])

    doc.save(str(DEST))

    print(f"[OK] {DEST.name}")
    print(f"     Notas explicativas removidas      : {notas_count}")
    print(f"     Parágrafos FILL substituídos      : {fill_count}")
    print(f"     Placeholders inline substituídos  : {inline_count}")
    print(f"     Alternativas/OU removidos         : {alt_count}")
    print("Template TR pronto.")


if __name__ == "__main__":
    main()
